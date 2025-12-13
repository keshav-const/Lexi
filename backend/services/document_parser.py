"""
Document Parser Service for DOCX and PDF files.
CREATED BY UOIONHHC
"""
import io
from docx import Document
import fitz  # PyMuPDF
from typing import Optional


async def parse_docx(file_content: bytes) -> str:
    """
    Parse a DOCX file and extract text content.
    
    Args:
        file_content: Raw bytes of the DOCX file
    
    Returns:
        Extracted text content
    """
    doc = Document(io.BytesIO(file_content))
    
    paragraphs = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    return "\n\n".join(paragraphs)


async def parse_pdf(file_content: bytes) -> str:
    """
    Parse a PDF file and extract text content.
    
    Args:
        file_content: Raw bytes of the PDF file
    
    Returns:
        Extracted text content
    """
    doc = fitz.open(stream=file_content, filetype="pdf")
    
    text_parts = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text = page.get_text()
        if text.strip():
            text_parts.append(text.strip())
    
    doc.close()
    return "\n\n".join(text_parts)


async def parse_document(file_content: bytes, filename: str) -> tuple[str, str]:
    """
    Parse a document based on its file extension.
    
    Args:
        file_content: Raw bytes of the file
        filename: Original filename (used to determine type)
    
    Returns:
        Tuple of (extracted_text, mime_type)
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.docx'):
        text = await parse_docx(file_content)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif filename_lower.endswith('.pdf'):
        text = await parse_pdf(file_content)
        mime = "application/pdf"
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    
    return text, mime


def chunk_text(text: str, chunk_size: int = 8000, overlap: int = 500) -> list[str]:
    """
    Split text into overlapping chunks for processing.
    
    Args:
        text: Text to chunk
        chunk_size: Maximum chunk size in characters
        overlap: Number of overlapping characters between chunks
    
    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at a paragraph or sentence boundary
        if end < len(text):
            # Look for paragraph break
            para_break = text.rfind('\n\n', start, end)
            if para_break > start + chunk_size // 2:
                end = para_break + 2
            else:
                # Look for sentence break
                sentence_break = text.rfind('. ', start, end)
                if sentence_break > start + chunk_size // 2:
                    end = sentence_break + 2
        
        chunks.append(text[start:end].strip())
        start = end - overlap
    
    return chunks


def get_text_preview(text: str, max_length: int = 500) -> str:
    """Get a preview of the text, truncated if necessary."""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."
